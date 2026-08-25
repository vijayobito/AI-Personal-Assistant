"""
AI Personal Assistant & Productivity Manager — FastAPI Backend
Handles Authentication, Conversations, Messages, Memory, Tasks & Reminders, AI Notes,
AI Daily Planner, Multi-format Document RAG (PDF, DOCX, TXT, CSV), Universal Search,
Analytics, Settings, Notifications, and Autonomous AI Agent Execution.
"""

import os
import uuid
import json
import io
import re
from datetime import datetime, timezone, timedelta
from contextlib import asynccontextmanager

from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Depends, Query, Header, status
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, EmailStr
from sqlalchemy import create_engine, Column, String, Text, DateTime, ForeignKey, Integer, Boolean
from sqlalchemy.orm import declarative_base, sessionmaker, relationship

import jwt
from passlib.context import CryptContext

from pypdf import PdfReader
import docx
import pandas as pd

from services.ai_service import generate_response_with_tools, generate_embedding, cosine_similarity

# ──────────────────────────────────────────────
# Security & JWT Setup
# ──────────────────────────────────────────────

SECRET_KEY = os.getenv("JWT_SECRET_KEY", "nexus_ai_super_secret_jwt_key_2026_change_in_prod")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_DAYS = 30

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

def hash_password(password: str) -> str:
    return pwd_context.hash(password)

def verify_password(plain_password: str, hashed_password: str) -> bool:
    return pwd_context.verify(plain_password, hashed_password)

def create_access_token(user_id: str) -> str:
    expire = datetime.now(timezone.utc) + timedelta(days=ACCESS_TOKEN_EXPIRE_DAYS)
    payload = {"sub": user_id, "exp": expire}
    return jwt.encode(payload, SECRET_KEY, algorithm=ALGORITHM)

# ──────────────────────────────────────────────
# Database Setup
# ──────────────────────────────────────────────

DATABASE_URL = "sqlite:///./chat_history.db"
engine = create_engine(DATABASE_URL, connect_args={"check_same_thread": False})
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()


# ──────────────────────────────────────────────
# ORM Models
# ──────────────────────────────────────────────

class User(Base):
    __tablename__ = "users"

    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    name = Column(String, nullable=False)
    email = Column(String, unique=True, nullable=False, index=True)
    hashed_password = Column(String, nullable=False)
    profile_image = Column(String, default="")
    account_plan = Column(String, default="Pro Plan")
    preferences_json = Column(Text, default="{}")
    settings_json = Column(Text, default="{}")
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))

    conversations = relationship("Conversation", back_populates="user", cascade="all, delete-orphan")
    reminders = relationship("Reminder", back_populates="user", cascade="all, delete-orphan")
    tasks = relationship("Task", back_populates="user", cascade="all, delete-orphan")
    memories = relationship("Memory", back_populates="user", cascade="all, delete-orphan")
    notes = relationship("Note", back_populates="user", cascade="all, delete-orphan")
    daily_plans = relationship("DailyPlan", back_populates="user", cascade="all, delete-orphan")
    documents = relationship("Document", back_populates="user", cascade="all, delete-orphan")
    notifications = relationship("Notification", back_populates="user", cascade="all, delete-orphan")


class Conversation(Base):
    __tablename__ = "conversations"

    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    user_id = Column(String, ForeignKey("users.id"), nullable=True)
    title = Column(String, default="New Chat")
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    updated_at = Column(DateTime, default=lambda: datetime.now(timezone.utc),
                        onupdate=lambda: datetime.now(timezone.utc))

    user = relationship("User", back_populates="conversations")
    messages = relationship("Message", back_populates="conversation",
                            cascade="all, delete-orphan",
                            order_by="Message.created_at")


class Message(Base):
    __tablename__ = "messages"

    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    conversation_id = Column(String, ForeignKey("conversations.id"), nullable=False)
    role = Column(String, nullable=False)  # "user" or "assistant"
    content = Column(Text, nullable=False)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))

    conversation = relationship("Conversation", back_populates="messages")


class Reminder(Base):
    __tablename__ = "reminders"

    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    user_id = Column(String, ForeignKey("users.id"), nullable=True)
    task = Column(String, nullable=False)
    date = Column(String, nullable=False)
    time = Column(String, nullable=False)
    priority = Column(String, default="Medium")
    recurrence = Column(String, default="none")
    status = Column(String, default="pending")  # "pending", "completed"
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))

    user = relationship("User", back_populates="reminders")


class Task(Base):
    __tablename__ = "tasks"

    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    user_id = Column(String, ForeignKey("users.id"), nullable=True)
    title = Column(String, nullable=False)
    description = Column(Text, default="")
    priority = Column(String, default="Medium")  # "High", "Medium", "Low"
    due_date = Column(String, default="Today")
    due_time = Column(String, default="7:00 PM")
    category = Column(String, default="General")
    recurrence = Column(String, default="none")  # "none", "daily", "weekly"
    status = Column(String, default="pending")  # "pending", "completed"
    completed_at = Column(DateTime, nullable=True)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))

    user = relationship("User", back_populates="tasks")


class Memory(Base):
    __tablename__ = "memories"

    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    user_id = Column(String, ForeignKey("users.id"), nullable=True)
    fact = Column(Text, nullable=False)
    category = Column(String, default="general")
    importance = Column(String, default="Medium")
    is_active = Column(Boolean, default=True)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    updated_at = Column(DateTime, default=lambda: datetime.now(timezone.utc),
                        onupdate=lambda: datetime.now(timezone.utc))

    user = relationship("User", back_populates="memories")


class Note(Base):
    __tablename__ = "notes"

    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    user_id = Column(String, ForeignKey("users.id"), nullable=True)
    title = Column(String, nullable=False)
    content = Column(Text, nullable=False)
    tags = Column(String, default="")
    is_pinned = Column(Boolean, default=False)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    updated_at = Column(DateTime, default=lambda: datetime.now(timezone.utc),
                        onupdate=lambda: datetime.now(timezone.utc))

    user = relationship("User", back_populates="notes")


class DailyPlan(Base):
    __tablename__ = "daily_plans"

    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    user_id = Column(String, ForeignKey("users.id"), nullable=True)
    date = Column(String, nullable=False)
    schedule_json = Column(Text, nullable=False)  # JSON array of schedule slots
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))

    user = relationship("User", back_populates="daily_plans")


class Document(Base):
    __tablename__ = "documents"

    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    user_id = Column(String, ForeignKey("users.id"), nullable=True)
    filename = Column(String, nullable=False)
    file_type = Column(String, default="pdf")  # pdf, docx, txt, csv
    text_content = Column(Text, nullable=False)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))

    user = relationship("User", back_populates="documents")
    chunks = relationship("DocumentChunk", back_populates="document", cascade="all, delete-orphan")


class DocumentChunk(Base):
    __tablename__ = "document_chunks"

    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    document_id = Column(String, ForeignKey("documents.id"), nullable=False)
    chunk_index = Column(Integer, nullable=False)
    content = Column(Text, nullable=False)
    embedding_json = Column(Text, nullable=False)

    document = relationship("Document", back_populates="chunks")


class Notification(Base):
    __tablename__ = "notifications"

    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    user_id = Column(String, ForeignKey("users.id"), nullable=True)
    title = Column(String, nullable=False)
    message = Column(String, nullable=False)
    type = Column(String, default="info")  # info, reminder, task, document
    is_read = Column(Boolean, default=False)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))

    user = relationship("User", back_populates="notifications")


# ──────────────────────────────────────────────
# Pydantic Schemas
# ──────────────────────────────────────────────

class UserRegister(BaseModel):
    name: str
    email: EmailStr
    password: str


class UserLogin(BaseModel):
    email: EmailStr
    password: str


class UserUpdate(BaseModel):
    name: str | None = None
    email: EmailStr | None = None
    profile_image: str | None = None
    account_plan: str | None = None
    preferences: dict | None = None
    settings: dict | None = None


class UserOut(BaseModel):
    id: str
    name: str
    email: str
    profile_image: str
    account_plan: str
    preferences: dict = {}
    settings: dict = {}
    created_at: datetime

    class Config:
        from_attributes = True


class AuthResponse(BaseModel):
    token: str
    user: UserOut


class ChatRequest(BaseModel):
    message: str
    conversation_id: str | None = None


class MessageOut(BaseModel):
    id: str
    role: str
    content: str
    created_at: datetime

    class Config:
        from_attributes = True


class ConversationOut(BaseModel):
    id: str
    title: str
    created_at: datetime
    updated_at: datetime

    class Config:
        from_attributes = True


class ConversationDetail(ConversationOut):
    messages: list[MessageOut] = []


class ChatResponse(BaseModel):
    conversation_id: str
    user_message: MessageOut
    ai_message: MessageOut


class ReminderCreate(BaseModel):
    task: str
    date: str
    time: str
    priority: str = "Medium"
    recurrence: str = "none"


class ReminderOut(BaseModel):
    id: str
    task: str
    date: str
    time: str
    priority: str = "Medium"
    recurrence: str = "none"
    status: str
    created_at: datetime

    class Config:
        from_attributes = True


class TaskCreate(BaseModel):
    title: str
    description: str = ""
    priority: str = "Medium"
    due_date: str = "Today"
    due_time: str = "7:00 PM"
    category: str = "General"
    recurrence: str = "none"


class TaskOut(BaseModel):
    id: str
    title: str
    description: str
    priority: str
    due_date: str
    due_time: str
    category: str = "General"
    recurrence: str
    status: str
    created_at: datetime

    class Config:
        from_attributes = True


class MemoryCreate(BaseModel):
    fact: str
    category: str = "general"
    importance: str = "Medium"


class MemoryOut(BaseModel):
    id: str
    fact: str
    category: str
    importance: str = "Medium"
    is_active: bool
    created_at: datetime

    class Config:
        from_attributes = True


class NoteCreate(BaseModel):
    title: str
    content: str
    tags: str = ""
    is_pinned: bool = False


class NoteOut(BaseModel):
    id: str
    title: str
    content: str
    tags: str
    is_pinned: bool
    created_at: datetime
    updated_at: datetime

    class Config:
        from_attributes = True


class PlanGenerateRequest(BaseModel):
    activities: list[str] = []
    date: str = "Today"


class DailyPlanOut(BaseModel):
    id: str
    date: str
    schedule: list[dict]
    created_at: datetime


class DocumentOut(BaseModel):
    id: str
    filename: str
    file_type: str
    created_at: datetime
    chunk_count: int = 0

    class Config:
        from_attributes = True


class NotificationOut(BaseModel):
    id: str
    title: str
    message: str
    type: str
    is_read: bool
    created_at: datetime

    class Config:
        from_attributes = True


class CommandCenterRequest(BaseModel):
    command: str


# ──────────────────────────────────────────────
# App Lifespan & Helpers
# ──────────────────────────────────────────────

@asynccontextmanager
async def lifespan(app: FastAPI):
    Base.metadata.create_all(bind=engine)
    db = SessionLocal()
    try:
        if db.query(User).count() == 0:
            demo_user = User(
                id="default_user_1",
                name="Vijay Kumar",
                email="vijay@nexus.ai",
                hashed_password=hash_password("password123"),
                profile_image="",
                account_plan="Pro Plan",
                preferences_json=json.dumps({"ai_model": "Gemini 2.0 Flash", "memory_system": True}),
                settings_json=json.dumps({"theme": "dark", "notifications": True, "voice_input": True})
            )
            db.add(demo_user)
            db.commit()
            print("[OK] Created default demo user: vijay@nexus.ai / password123")
    finally:
        db.close()

    print("[OK] Database tables created & verified successfully")
    yield


app = FastAPI(
    title="AI Personal Assistant & Productivity Manager",
    version="3.0.0",
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


def get_current_user(authorization: str = Header(None), db=Depends(get_db)) -> User:
    if authorization and authorization.startswith("Bearer "):
        token = authorization.split(" ")[1]
        try:
            payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
            user_id = payload.get("sub")
            if user_id:
                user = db.query(User).filter(User.id == user_id).first()
                if user:
                    return user
        except jwt.PyJWTError:
            pass

    user = db.query(User).first()
    if not user:
        user = User(
            name="Nexus User",
            email="user@nexus.ai",
            hashed_password=hash_password("password123"),
            account_plan="Pro Plan"
        )
        db.add(user)
        db.commit()
        db.refresh(user)
    return user


def _auto_title(text: str, max_len: int = 40) -> str:
    clean = text.strip().replace("\n", " ")
    return clean[:max_len] + ("…" if len(clean) > max_len else "")


def _parse_user_json(user: User):
    pref = json.loads(user.preferences_json) if user.preferences_json else {}
    sett = json.loads(user.settings_json) if user.settings_json else {}
    return UserOut(
        id=user.id,
        name=user.name,
        email=user.email,
        profile_image=user.profile_image or "",
        account_plan=user.account_plan or "Pro Plan",
        preferences=pref,
        settings=sett,
        created_at=user.created_at
    )


# ──────────────────────────────────────────────
# Autonomous Tool Execution Dispatcher
# ──────────────────────────────────────────────

def execute_ai_tool(tool_name: str, tool_args: dict, user_id: str | None = None) -> dict:
    """Executes backend data operations invoked autonomously by Gemini tools."""
    db = SessionLocal()
    try:
        if tool_name == "remember_user_fact":
            fact = tool_args.get("fact", "").strip()
            cat = tool_args.get("category", "general").strip()
            if fact:
                mem = Memory(user_id=user_id, fact=fact, category=cat)
                db.add(mem)
                db.commit()
                db.refresh(mem)
                return {"success": True, "message": f"Remembered: '{fact}'", "memory_id": mem.id}
            return {"success": False, "message": "Fact cannot be empty"}

        elif tool_name == "create_task":
            title = tool_args.get("title", "New Task")
            priority = tool_args.get("priority", "Medium")
            due_date = tool_args.get("due_date", "Today")
            due_time = tool_args.get("due_time", "7:00 PM")
            recurrence = tool_args.get("recurrence", "none")

            tsk = Task(user_id=user_id, title=title, priority=priority, due_date=due_date, due_time=due_time, recurrence=recurrence)
            db.add(tsk)
            db.commit()
            db.refresh(tsk)
            return {"success": True, "message": f"Created task '{title}' ({priority} priority)", "task_id": tsk.id}

        elif tool_name == "create_reminder":
            task = tool_args.get("task", "Reminder Task")
            date = tool_args.get("date", "Tomorrow")
            time_val = tool_args.get("time", "7:00 PM")

            rem = Reminder(user_id=user_id, task=task, date=date, time=time_val)
            db.add(rem)
            db.commit()
            db.refresh(rem)
            return {"success": True, "message": f"Created reminder '{task}' for {date} at {time_val}", "reminder_id": rem.id}

        elif tool_name == "get_reminders":
            q = db.query(Reminder)
            if user_id:
                q = q.filter(Reminder.user_id == user_id)
            reminders = q.order_by(Reminder.created_at.desc()).all()
            return {"reminders": [{"id": r.id, "task": r.task, "date": r.date, "time": r.time, "status": r.status} for r in reminders]}

        elif tool_name == "delete_reminder":
            rem_id = tool_args.get("reminder_id")
            task_search = tool_args.get("task_search")
            q = db.query(Reminder)
            if user_id:
                q = q.filter(Reminder.user_id == user_id)
            rem = None
            if rem_id:
                rem = q.filter(Reminder.id == rem_id).first()
            elif task_search:
                rem = q.filter(Reminder.task.ilike(f"%{task_search}%")).first()

            if rem:
                title = rem.task
                db.delete(rem)
                db.commit()
                return {"success": True, "message": f"Deleted reminder '{title}'"}
            return {"success": False, "message": "Reminder not found"}

        elif tool_name == "create_note":
            title = tool_args.get("title", "Untitled Note")
            content = tool_args.get("content", "")
            tags = tool_args.get("tags", "")

            nt = Note(user_id=user_id, title=title, content=content, tags=tags)
            db.add(nt)
            db.commit()
            db.refresh(nt)
            return {"success": True, "message": f"Saved note '{title}'", "note_id": nt.id}

        elif tool_name == "get_notes":
            query = tool_args.get("query", "").lower()
            q = db.query(Note)
            if user_id:
                q = q.filter(Note.user_id == user_id)
            if query:
                q = q.filter(Note.title.ilike(f"%{query}%") | Note.content.ilike(f"%{query}%") | Note.tags.ilike(f"%{query}%"))
            notes = q.order_by(Note.is_pinned.desc(), Note.updated_at.desc()).all()
            return {"notes": [{"id": n.id, "title": n.title, "content": n.content, "tags": n.tags} for n in notes]}

        elif tool_name == "generate_daily_plan":
            activities = tool_args.get("activities", [])
            schedule = []
            default_slots = [
                ("9:00 AM", "Morning"), ("11:00 AM", "Morning"),
                ("2:00 PM", "Afternoon"), ("4:00 PM", "Afternoon"),
                ("7:00 PM", "Evening"), ("9:00 PM", "Evening")
            ]
            for idx, act in enumerate(activities):
                t_str, period = default_slots[idx % len(default_slots)]
                schedule.append({"id": str(uuid.uuid4()), "time": t_str, "period": period, "activity": act, "status": "pending"})

            if not schedule:
                schedule = [
                    {"id": str(uuid.uuid4()), "time": "9:00 AM", "period": "Morning", "activity": "Morning Review & Focused Study", "status": "pending"},
                    {"id": str(uuid.uuid4()), "time": "11:30 AM", "period": "Morning", "activity": "Core Project Development", "status": "pending"},
                    {"id": str(uuid.uuid4()), "time": "3:00 PM", "period": "Afternoon", "activity": "Deep Focus Work & Task Execution", "status": "pending"},
                    {"id": str(uuid.uuid4()), "time": "6:00 PM", "period": "Evening", "activity": "Exercise & Refresh", "status": "pending"},
                    {"id": str(uuid.uuid4()), "time": "8:00 PM", "period": "Evening", "activity": "Evening Plan Review & Brainstorming", "status": "pending"}
                ]

            plan = DailyPlan(user_id=user_id, date="Today", schedule_json=json.dumps(schedule))
            db.add(plan)
            db.commit()
            db.refresh(plan)
            return {"success": True, "schedule": schedule, "plan_id": plan.id}

        elif tool_name == "search_all":
            sq = tool_args.get("query", "").strip().lower()
            t_q = db.query(Task)
            n_q = db.query(Note)
            m_q = db.query(Memory)
            d_q = db.query(Document)
            if user_id:
                t_q = t_q.filter(Task.user_id == user_id)
                n_q = n_q.filter(Note.user_id == user_id)
                m_q = m_q.filter(Memory.user_id == user_id)
                d_q = d_q.filter(Document.user_id == user_id)

            tasks = t_q.filter(Task.title.ilike(f"%{sq}%") | Task.description.ilike(f"%{sq}%")).all()
            notes = n_q.filter(Note.title.ilike(f"%{sq}%") | Note.content.ilike(f"%{sq}%")).all()
            memories = m_q.filter(Memory.fact.ilike(f"%{sq}%")).all()
            docs = d_q.filter(Document.filename.ilike(f"%{sq}%") | Document.text_content.ilike(f"%{sq}%")).all()

            return {
                "tasks": [{"id": t.id, "title": t.title, "priority": t.priority} for t in tasks],
                "notes": [{"id": n.id, "title": n.title, "content": n.content[:100]} for n in notes],
                "memories": [{"id": m.id, "fact": m.fact} for m in memories],
                "documents": [{"id": d.id, "filename": d.filename} for d in docs]
            }

        elif tool_name == "search_documents":
            search_query = tool_args.get("query", "")
            query_emb = generate_embedding(search_query)
            q = db.query(DocumentChunk).join(Document)
            if user_id:
                q = q.filter(Document.user_id == user_id)
            chunks = q.all()
            scored = []
            for c in chunks:
                emb = json.loads(c.embedding_json)
                score = cosine_similarity(query_emb, emb)
                scored.append((score, c.content))
            scored.sort(key=lambda x: x[0], reverse=True)
            top_passages = [content for score, content in scored[:3]]
            return {
                "query": search_query,
                "results": "\n\n".join(top_passages) if top_passages else "No matching document content found."
            }

        elif tool_name == "summarize_document":
            doc_id = tool_args.get("document_id")
            q = db.query(Document)
            if user_id:
                q = q.filter(Document.user_id == user_id)
            if doc_id:
                doc = q.filter(Document.id == doc_id).first()
            else:
                doc = q.order_by(Document.created_at.desc()).first()

            if not doc:
                return {"summary": "No uploaded document found to summarize."}
            return {
                "filename": doc.filename,
                "summary": f"Document: {doc.filename}\n\nExcerpt:\n{doc.text_content[:1500]}..."
            }

        return {"error": f"Unknown tool: {tool_name}"}
    finally:
        db.close()


# ──────────────────────────────────────────────
# Routes — Authentication & User Profile
# ──────────────────────────────────────────────

@app.post("/api/auth/register", response_model=AuthResponse)
def register(req: UserRegister, db=Depends(get_db)):
    existing = db.query(User).filter(User.email == req.email.lower()).first()
    if existing:
        raise HTTPException(status_code=400, detail="User with this email already exists")

    if len(req.password) < 6:
        raise HTTPException(status_code=400, detail="Password must be at least 6 characters long")

    user = User(
        name=req.name,
        email=req.email.lower(),
        hashed_password=hash_password(req.password),
        account_plan="Pro Plan",
        preferences_json=json.dumps({"ai_model": "Gemini 2.0 Flash", "memory_system": True}),
        settings_json=json.dumps({"theme": "dark", "notifications": True, "voice_input": True})
    )
    db.add(user)
    db.commit()
    db.refresh(user)

    token = create_access_token(user.id)
    return AuthResponse(token=token, user=_parse_user_json(user))


@app.post("/api/auth/login", response_model=AuthResponse)
def login(req: UserLogin, db=Depends(get_db)):
    user = db.query(User).filter(User.email == req.email.lower()).first()
    if not user or not verify_password(req.password, user.hashed_password):
        raise HTTPException(status_code=401, detail="Invalid email or password")

    token = create_access_token(user.id)
    return AuthResponse(token=token, user=_parse_user_json(user))


@app.get("/api/auth/me", response_model=UserOut)
def get_current_user_profile(user: User = Depends(get_current_user)):
    return _parse_user_json(user)


@app.put("/api/auth/profile", response_model=UserOut)
def update_profile(req: UserUpdate, user: User = Depends(get_current_user), db=Depends(get_db)):
    if req.name is not None:
        user.name = req.name
    if req.email is not None:
        user.email = req.email.lower()
    if req.profile_image is not None:
        user.profile_image = req.profile_image
    if req.account_plan is not None:
        user.account_plan = req.account_plan

    if req.preferences is not None:
        curr_pref = json.loads(user.preferences_json) if user.preferences_json else {}
        curr_pref.update(req.preferences)
        user.preferences_json = json.dumps(curr_pref)

    if req.settings is not None:
        curr_sett = json.loads(user.settings_json) if user.settings_json else {}
        curr_sett.update(req.settings)
        user.settings_json = json.dumps(curr_sett)

    db.commit()
    db.refresh(user)
    return _parse_user_json(user)


# ──────────────────────────────────────────────
# Routes — AI Chat Workspace
# ──────────────────────────────────────────────

@app.post("/api/conversations", response_model=ConversationOut)
def create_conversation(user: User = Depends(get_current_user), db=Depends(get_db)):
    conv = Conversation(user_id=user.id)
    db.add(conv)
    db.commit()
    db.refresh(conv)
    return conv


@app.get("/api/conversations", response_model=list[ConversationOut])
def list_conversations(user: User = Depends(get_current_user), db=Depends(get_db)):
    return db.query(Conversation).filter(Conversation.user_id == user.id).order_by(Conversation.updated_at.desc()).all()


@app.get("/api/conversations/{conversation_id}", response_model=ConversationDetail)
def get_conversation(conversation_id: str, user: User = Depends(get_current_user), db=Depends(get_db)):
    conv = db.query(Conversation).filter(Conversation.id == conversation_id, Conversation.user_id == user.id).first()
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")
    return conv


@app.delete("/api/conversations/{conversation_id}")
def delete_conversation(conversation_id: str, user: User = Depends(get_current_user), db=Depends(get_db)):
    conv = db.query(Conversation).filter(Conversation.id == conversation_id, Conversation.user_id == user.id).first()
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")
    db.delete(conv)
    db.commit()
    return {"detail": "Conversation deleted"}


@app.post("/api/chat", response_model=ChatResponse)
def chat(req: ChatRequest, user: User = Depends(get_current_user), db=Depends(get_db)):
    if req.conversation_id:
        conv = db.query(Conversation).filter(Conversation.id == req.conversation_id, Conversation.user_id == user.id).first()
        if not conv:
            raise HTTPException(status_code=404, detail="Conversation not found")
    else:
        conv = Conversation(user_id=user.id, title=_auto_title(req.message))
        db.add(conv)
        db.commit()
        db.refresh(conv)

    user_msg = Message(
        conversation_id=conv.id,
        role="user",
        content=req.message,
    )
    db.add(user_msg)
    db.commit()
    db.refresh(user_msg)

    rem_match = re.search(r"remind me to (.+?) (tomorrow|today|\d{4}-\d{2}-\d{2}) at (\d{1,2}(?::\d{2})?\s*(?:am|pm|AM|PM))", req.message, re.IGNORECASE)
    if rem_match:
        task_str = rem_match.group(1).strip()
        date_str = rem_match.group(2).capitalize()
        time_str = rem_match.group(3).upper()
        rem = Reminder(user_id=user.id, task=task_str, date=date_str, time=time_str)
        db.add(rem)
        db.commit()

    active_mems = db.query(Memory).filter(Memory.user_id == user.id, Memory.is_active == True).all()
    mem_context = ""
    if active_mems:
        mem_str = "; ".join([m.fact for m in active_mems])
        mem_context = f"[Active Saved Memories: {mem_str}]\n"

    history = []
    if mem_context:
        history.append({"role": "user", "content": mem_context})
        history.append({"role": "assistant", "content": "Understood. I will keep your saved facts in mind."})

    for m in conv.messages:
        history.append({"role": m.role, "content": m.content})

    def bound_tool_executor(tool_name, tool_args):
        return execute_ai_tool(tool_name, tool_args, user_id=user.id)

    ai_text = generate_response_with_tools(history, tool_executor_func=bound_tool_executor)

    ai_msg = Message(
        conversation_id=conv.id,
        role="assistant",
        content=ai_text,
    )
    db.add(ai_msg)

    if len(conv.messages) <= 2:
        conv.title = _auto_title(req.message)

    conv.updated_at = datetime.now(timezone.utc)
    db.commit()
    db.refresh(ai_msg)

    return ChatResponse(
        conversation_id=conv.id,
        user_message=MessageOut.model_validate(user_msg),
        ai_message=MessageOut.model_validate(ai_msg),
    )


# ──────────────────────────────────────────────
# Routes — Smart AI Memory
# ──────────────────────────────────────────────

@app.get("/api/memories", response_model=list[MemoryOut])
def get_memories(user: User = Depends(get_current_user), db=Depends(get_db)):
    return db.query(Memory).filter(Memory.user_id == user.id).order_by(Memory.created_at.desc()).all()


@app.post("/api/memories", response_model=MemoryOut)
def create_memory(req: MemoryCreate, user: User = Depends(get_current_user), db=Depends(get_db)):
    mem = Memory(user_id=user.id, fact=req.fact, category=req.category, importance=req.importance)
    db.add(mem)
    db.commit()
    db.refresh(mem)
    return mem


@app.patch("/api/memories/{memory_id}/toggle")
def toggle_memory(memory_id: str, user: User = Depends(get_current_user), db=Depends(get_db)):
    mem = db.query(Memory).filter(Memory.id == memory_id, Memory.user_id == user.id).first()
    if not mem:
        raise HTTPException(status_code=404, detail="Memory not found")
    mem.is_active = not mem.is_active
    db.commit()
    return {"detail": "Toggled memory status", "is_active": mem.is_active}


@app.delete("/api/memories/{memory_id}")
def delete_memory(memory_id: str, user: User = Depends(get_current_user), db=Depends(get_db)):
    mem = db.query(Memory).filter(Memory.id == memory_id, Memory.user_id == user.id).first()
    if not mem:
        raise HTTPException(status_code=404, detail="Memory not found")
    db.delete(mem)
    db.commit()
    return {"detail": "Memory deleted"}


# ──────────────────────────────────────────────
# Routes — Smart Tasks & Reminders Manager
# ──────────────────────────────────────────────

@app.get("/api/tasks", response_model=list[TaskOut])
def get_tasks(user: User = Depends(get_current_user), db=Depends(get_db)):
    return db.query(Task).filter(Task.user_id == user.id).order_by(Task.created_at.desc()).all()


@app.post("/api/tasks", response_model=TaskOut)
def create_task(req: TaskCreate, user: User = Depends(get_current_user), db=Depends(get_db)):
    tsk = Task(
        user_id=user.id,
        title=req.title,
        description=req.description,
        priority=req.priority,
        due_date=req.due_date,
        due_time=req.due_time,
        category=req.category,
        recurrence=req.recurrence
    )
    db.add(tsk)
    db.commit()
    db.refresh(tsk)
    return tsk


@app.put("/api/tasks/{task_id}", response_model=TaskOut)
def update_task(task_id: str, req: TaskCreate, user: User = Depends(get_current_user), db=Depends(get_db)):
    tsk = db.query(Task).filter(Task.id == task_id, Task.user_id == user.id).first()
    if not tsk:
        raise HTTPException(status_code=404, detail="Task not found")
    tsk.title = req.title
    tsk.description = req.description
    tsk.priority = req.priority
    tsk.due_date = req.due_date
    tsk.due_time = req.due_time
    tsk.category = req.category
    tsk.recurrence = req.recurrence
    db.commit()
    db.refresh(tsk)
    return tsk


@app.patch("/api/tasks/{task_id}/complete")
def complete_task(task_id: str, user: User = Depends(get_current_user), db=Depends(get_db)):
    tsk = db.query(Task).filter(Task.id == task_id, Task.user_id == user.id).first()
    if not tsk:
        raise HTTPException(status_code=404, detail="Task not found")
    if tsk.status == "pending":
        tsk.status = "completed"
        tsk.completed_at = datetime.now(timezone.utc)
    else:
        tsk.status = "pending"
        tsk.completed_at = None
    db.commit()
    return {"detail": "Updated task status", "status": tsk.status}


@app.delete("/api/tasks/{task_id}")
def delete_task(task_id: str, user: User = Depends(get_current_user), db=Depends(get_db)):
    tsk = db.query(Task).filter(Task.id == task_id, Task.user_id == user.id).first()
    if not tsk:
        raise HTTPException(status_code=404, detail="Task not found")
    db.delete(tsk)
    db.commit()
    return {"detail": "Task deleted"}


@app.post("/api/reminders", response_model=ReminderOut)
def create_reminder(req: ReminderCreate, user: User = Depends(get_current_user), db=Depends(get_db)):
    rem = Reminder(
        user_id=user.id,
        task=req.task,
        date=req.date,
        time=req.time,
        priority=req.priority,
        recurrence=req.recurrence
    )
    db.add(rem)
    db.commit()
    db.refresh(rem)
    return rem


@app.get("/api/reminders", response_model=list[ReminderOut])
def get_reminders(user: User = Depends(get_current_user), db=Depends(get_db)):
    return db.query(Reminder).filter(Reminder.user_id == user.id).order_by(Reminder.created_at.desc()).all()


@app.patch("/api/reminders/{reminder_id}/complete")
def complete_reminder(reminder_id: str, user: User = Depends(get_current_user), db=Depends(get_db)):
    rem = db.query(Reminder).filter(Reminder.id == reminder_id, Reminder.user_id == user.id).first()
    if not rem:
        raise HTTPException(status_code=404, detail="Reminder not found")
    rem.status = "completed" if rem.status == "pending" else "pending"
    db.commit()
    return {"detail": "Updated reminder status", "status": rem.status}


@app.delete("/api/reminders/{reminder_id}")
def delete_reminder(reminder_id: str, user: User = Depends(get_current_user), db=Depends(get_db)):
    rem = db.query(Reminder).filter(Reminder.id == reminder_id, Reminder.user_id == user.id).first()
    if not rem:
        raise HTTPException(status_code=404, detail="Reminder not found")
    db.delete(rem)
    db.commit()
    return {"detail": "Reminder deleted"}


# ──────────────────────────────────────────────
# Routes — AI Daily Planner
# ──────────────────────────────────────────────

@app.post("/api/planner/generate", response_model=DailyPlanOut)
def generate_daily_plan_route(req: PlanGenerateRequest, user: User = Depends(get_current_user), db=Depends(get_db)):
    if not req.activities:
        pending_tasks = db.query(Task).filter(Task.user_id == user.id, Task.status == "pending").all()
        req.activities = [t.title for t in pending_tasks] if pending_tasks else ["Study Python", "Work on AI Project", "Exercise"]

    result = execute_ai_tool("generate_daily_plan", {"activities": req.activities}, user_id=user.id)
    schedule = result.get("schedule", [])
    
    plan = db.query(DailyPlan).filter(DailyPlan.id == result.get("plan_id")).first()
    if not plan:
        plan = DailyPlan(user_id=user.id, date=req.date, schedule_json=json.dumps(schedule))
        db.add(plan)
        db.commit()
        db.refresh(plan)

    return DailyPlanOut(
        id=plan.id,
        date=plan.date,
        schedule=schedule,
        created_at=plan.created_at
    )


@app.get("/api/planner/today", response_model=DailyPlanOut)
def get_today_plan(user: User = Depends(get_current_user), db=Depends(get_db)):
    plan = db.query(DailyPlan).filter(DailyPlan.user_id == user.id).order_by(DailyPlan.created_at.desc()).first()
    if not plan:
        default_schedule = [
            {"id": str(uuid.uuid4()), "time": "9:00 AM", "period": "Morning", "activity": "Study Python", "status": "pending"},
            {"id": str(uuid.uuid4()), "time": "11:00 AM", "period": "Morning", "activity": "Work on AI Project", "status": "pending"},
            {"id": str(uuid.uuid4()), "time": "3:00 PM", "period": "Afternoon", "activity": "Focus Work & Code Review", "status": "pending"},
            {"id": str(uuid.uuid4()), "time": "5:00 PM", "period": "Evening", "activity": "Exercise", "status": "pending"},
            {"id": str(uuid.uuid4()), "time": "7:00 PM", "period": "Evening", "activity": "Review Progress", "status": "pending"}
        ]
        plan = DailyPlan(user_id=user.id, date="Today", schedule_json=json.dumps(default_schedule))
        db.add(plan)
        db.commit()
        db.refresh(plan)

    return DailyPlanOut(
        id=plan.id,
        date=plan.date,
        schedule=json.loads(plan.schedule_json),
        created_at=plan.created_at
    )


@app.put("/api/planner/{plan_id}/slots")
def update_plan_slots(plan_id: str, slots: list[dict], user: User = Depends(get_current_user), db=Depends(get_db)):
    plan = db.query(DailyPlan).filter(DailyPlan.id == plan_id, DailyPlan.user_id == user.id).first()
    if not plan:
        raise HTTPException(status_code=404, detail="Daily plan not found")
    plan.schedule_json = json.dumps(slots)
    db.commit()
    return {"detail": "Plan slots updated", "schedule": slots}


# ──────────────────────────────────────────────
# Routes — AI Notes
# ──────────────────────────────────────────────

@app.get("/api/notes", response_model=list[NoteOut])
def get_notes(query: str | None = None, user: User = Depends(get_current_user), db=Depends(get_db)):
    q = db.query(Note).filter(Note.user_id == user.id)
    if query:
        q = q.filter(Note.title.ilike(f"%{query}%") | Note.content.ilike(f"%{query}%") | Note.tags.ilike(f"%{query}%"))
    return q.order_by(Note.is_pinned.desc(), Note.updated_at.desc()).all()


@app.post("/api/notes", response_model=NoteOut)
def create_note(req: NoteCreate, user: User = Depends(get_current_user), db=Depends(get_db)):
    nt = Note(user_id=user.id, title=req.title, content=req.content, tags=req.tags, is_pinned=req.is_pinned)
    db.add(nt)
    db.commit()
    db.refresh(nt)
    return nt


@app.put("/api/notes/{note_id}", response_model=NoteOut)
def update_note(note_id: str, req: NoteCreate, user: User = Depends(get_current_user), db=Depends(get_db)):
    nt = db.query(Note).filter(Note.id == note_id, Note.user_id == user.id).first()
    if not nt:
        raise HTTPException(status_code=404, detail="Note not found")
    nt.title = req.title
    nt.content = req.content
    nt.tags = req.tags
    nt.is_pinned = req.is_pinned
    db.commit()
    db.refresh(nt)
    return nt


@app.patch("/api/notes/{note_id}/pin")
def toggle_pin_note(note_id: str, user: User = Depends(get_current_user), db=Depends(get_db)):
    nt = db.query(Note).filter(Note.id == note_id, Note.user_id == user.id).first()
    if not nt:
        raise HTTPException(status_code=404, detail="Note not found")
    nt.is_pinned = not nt.is_pinned
    db.commit()
    return {"detail": "Toggled pin status", "is_pinned": nt.is_pinned}


@app.delete("/api/notes/{note_id}")
def delete_note(note_id: str, user: User = Depends(get_current_user), db=Depends(get_db)):
    nt = db.query(Note).filter(Note.id == note_id, Note.user_id == user.id).first()
    if not nt:
        raise HTTPException(status_code=404, detail="Note not found")
    db.delete(nt)
    db.commit()
    return {"detail": "Note deleted"}


@app.post("/api/notes/{note_id}/ai-action")
def ai_note_action(note_id: str, action: str = Query(...), user: User = Depends(get_current_user), db=Depends(get_db)):
    nt = db.query(Note).filter(Note.id == note_id, Note.user_id == user.id).first()
    if not nt:
        raise HTTPException(status_code=404, detail="Note not found")

    prompt = f"Perform '{action}' action on this note content:\n\nTitle: {nt.title}\nContent: {nt.content}"
    res = generate_response_with_tools([{"role": "user", "content": prompt}])
    return {"action": action, "result": res}


# ──────────────────────────────────────────────
# Routes — Multi-Format Document Assistant (PDF, DOCX, TXT, CSV)
# ──────────────────────────────────────────────

@app.post("/api/documents/upload", response_model=DocumentOut)
async def upload_document(file: UploadFile = File(...), user: User = Depends(get_current_user), db=Depends(get_db)):
    filename = file.filename
    ext = os.path.splitext(filename)[1].lower()
    
    allowed_exts = {".pdf": "pdf", ".docx": "docx", ".txt": "txt", ".csv": "csv"}
    if ext not in allowed_exts:
        raise HTTPException(status_code=400, detail=f"Unsupported file format '{ext}'. Allowed: .pdf, .docx, .txt, .csv")

    file_type = allowed_exts[ext]
    extracted_text = ""

    try:
        contents = await file.read()
        
        if file_type == "pdf":
            reader = PdfReader(io.BytesIO(contents))
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    extracted_text += t + "\n"
        
        elif file_type == "docx":
            doc_file = docx.Document(io.BytesIO(contents))
            extracted_text = "\n".join([p.text for p in doc_file.paragraphs if p.text])
        
        elif file_type == "txt":
            extracted_text = contents.decode("utf-8", errors="ignore")
        
        elif file_type == "csv":
            df = pd.read_csv(io.BytesIO(contents))
            extracted_text = df.to_string()

        if not extracted_text.strip():
            raise HTTPException(status_code=400, detail=f"Could not extract readable text from '{filename}'")

        doc = Document(user_id=user.id, filename=filename, file_type=file_type, text_content=extracted_text)
        db.add(doc)
        db.commit()
        db.refresh(doc)

        chunk_size = 500
        overlap = 50
        chunks_text = []
        for i in range(0, len(extracted_text), chunk_size - overlap):
            chunk = extracted_text[i : i + chunk_size]
            if len(chunk.strip()) > 20:
                chunks_text.append(chunk)

        chunk_count = 0
        for idx, chunk_str in enumerate(chunks_text):
            emb_vec = generate_embedding(chunk_str)
            d_chunk = DocumentChunk(
                document_id=doc.id,
                chunk_index=idx,
                content=chunk_str,
                embedding_json=json.dumps(emb_vec),
            )
            db.add(d_chunk)
            chunk_count += 1

        db.commit()

        notif = Notification(user_id=user.id, title="Document Uploaded", message=f"Processed '{filename}' successfully.", type="document")
        db.add(notif)
        db.commit()

        return DocumentOut(
            id=doc.id,
            filename=doc.filename,
            file_type=doc.file_type,
            created_at=doc.created_at,
            chunk_count=chunk_count
        )

    except Exception as e:
        print(f"[ERROR] Document processing failed for {filename}: {e}")
        raise HTTPException(status_code=500, detail=f"Document parsing error: {str(e)}")


@app.get("/api/documents", response_model=list[DocumentOut])
def list_documents(user: User = Depends(get_current_user), db=Depends(get_db)):
    docs = db.query(Document).filter(Document.user_id == user.id).order_by(Document.created_at.desc()).all()
    res = []
    for d in docs:
        res.append(DocumentOut(
            id=d.id,
            filename=d.filename,
            file_type=d.file_type or "pdf",
            created_at=d.created_at,
            chunk_count=len(d.chunks)
        ))
    return res


@app.post("/api/documents/{document_id}/summarize")
def summarize_doc(document_id: str, user: User = Depends(get_current_user), db=Depends(get_db)):
    doc = db.query(Document).filter(Document.id == document_id, Document.user_id == user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    prompt = f"Provide a detailed, bulleted summary of this document:\n\nFilename: {doc.filename}\nContent Excerpt:\n{doc.text_content[:3000]}"
    res = generate_response_with_tools([{"role": "user", "content": prompt}])
    return {"filename": doc.filename, "summary": res}


@app.delete("/api/documents/{document_id}")
def delete_document(document_id: str, user: User = Depends(get_current_user), db=Depends(get_db)):
    doc = db.query(Document).filter(Document.id == document_id, Document.user_id == user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    db.delete(doc)
    db.commit()
    return {"detail": "Document deleted"}


# ──────────────────────────────────────────────
# Routes — Universal Search & Dashboard Analytics
# ──────────────────────────────────────────────

@app.get("/api/search")
def universal_search(q: str = Query(..., min_length=1), semantic: bool = False, user: User = Depends(get_current_user), db=Depends(get_db)):
    term = f"%{q.strip()}%"
    tasks = db.query(Task).filter(Task.user_id == user.id, Task.title.ilike(term) | Task.description.ilike(term)).all()
    reminders = db.query(Reminder).filter(Reminder.user_id == user.id, Reminder.task.ilike(term)).all()
    notes = db.query(Note).filter(Note.user_id == user.id, Note.title.ilike(term) | Note.content.ilike(term) | Note.tags.ilike(term)).all()
    memories = db.query(Memory).filter(Memory.user_id == user.id, Memory.fact.ilike(term)).all()
    documents = db.query(Document).filter(Document.user_id == user.id, Document.filename.ilike(term) | Document.text_content.ilike(term)).all()
    messages = db.query(Message).join(Conversation).filter(Conversation.user_id == user.id, Message.content.ilike(term)).limit(10).all()

    return {
        "query": q,
        "semantic": semantic,
        "results": {
            "tasks": [{"id": t.id, "title": t.title, "priority": t.priority, "status": t.status} for t in tasks],
            "reminders": [{"id": r.id, "task": r.task, "date": r.date, "time": r.time} for r in reminders],
            "notes": [{"id": n.id, "title": n.title, "content": n.content[:120], "tags": n.tags} for n in notes],
            "memories": [{"id": m.id, "fact": m.fact, "category": m.category} for m in memories],
            "documents": [{"id": d.id, "filename": d.filename, "file_type": d.file_type} for d in documents],
            "messages": [{"id": m.id, "role": m.role, "content": m.content[:150], "conversation_id": m.conversation_id} for m in messages]
        }
    }


@app.get("/api/analytics")
def get_analytics(user: User = Depends(get_current_user), db=Depends(get_db)):
    total_tasks = db.query(Task).filter(Task.user_id == user.id).count()
    completed_tasks = db.query(Task).filter(Task.user_id == user.id, Task.status == "completed").count()
    pending_tasks = total_tasks - completed_tasks
    completion_rate = int((completed_tasks / total_tasks * 100)) if total_tasks > 0 else 0

    total_reminders = db.query(Reminder).filter(Reminder.user_id == user.id).count()
    total_documents = db.query(Document).filter(Document.user_id == user.id).count()
    total_notes = db.query(Note).filter(Note.user_id == user.id).count()
    total_memories = db.query(Memory).filter(Memory.user_id == user.id).count()

    focus_time_hrs = round((completed_tasks * 45) / 60.0, 1) if completed_tasks > 0 else 0.0
    productivity_score = min(98, max(40, completion_rate + 15)) if total_tasks > 0 else 0

    # Build real weekly progress from completed_at timestamps (last 7 days)
    day_names = ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"]
    now = datetime.now(timezone.utc)
    weekly_progress = []
    completed_tasks_with_dates = db.query(Task).filter(
        Task.user_id == user.id,
        Task.status == "completed",
        Task.completed_at != None
    ).all()

    for offset in range(6, -1, -1):
        target_day = now - timedelta(days=offset)
        day_label = target_day.strftime("%a")  # Mon, Tue, etc.
        day_count = sum(
            1 for t in completed_tasks_with_dates
            if t.completed_at and t.completed_at.date() == target_day.date()
        )
        # Convert count to a 0-100 progress percentage (scale: 5+ tasks = 100%)
        progress_pct = min(100, day_count * 20) if day_count > 0 else 0
        weekly_progress.append({"day": day_label, "progress": progress_pct, "count": day_count})

    return {
        "total_tasks": total_tasks,
        "completed_tasks": completed_tasks,
        "pending_tasks": pending_tasks,
        "completion_rate": completion_rate,
        "productivity_score": productivity_score,
        "focus_time_hrs": focus_time_hrs,
        "total_reminders": total_reminders,
        "total_documents": total_documents,
        "total_notes": total_notes,
        "total_memories": total_memories,
        "weekly_progress": weekly_progress,
        "ai_insight": f"Your task completion rate is {completion_rate}%. You have {pending_tasks} pending tasks and {completed_tasks} completed. Keep up the momentum!"
    }


# ──────────────────────────────────────────────
# Routes — Notifications & Settings
# ──────────────────────────────────────────────

@app.get("/api/notifications", response_model=list[NotificationOut])
def get_notifications(user: User = Depends(get_current_user), db=Depends(get_db)):
    return db.query(Notification).filter(Notification.user_id == user.id).order_by(Notification.created_at.desc()).all()


@app.patch("/api/notifications/{notif_id}/read")
def mark_notification_read(notif_id: str, user: User = Depends(get_current_user), db=Depends(get_db)):
    notif = db.query(Notification).filter(Notification.id == notif_id, Notification.user_id == user.id).first()
    if notif:
        notif.is_read = True
        db.commit()
    return {"detail": "Marked read"}


@app.delete("/api/notifications/clear")
def clear_notifications(user: User = Depends(get_current_user), db=Depends(get_db)):
    db.query(Notification).filter(Notification.user_id == user.id).delete()
    db.commit()
    return {"detail": "Cleared all notifications"}


# ──────────────────────────────────────────────
# Routes — AI Command Center Execution
# ──────────────────────────────────────────────

@app.post("/api/command-center")
def execute_command_center(req: CommandCenterRequest, user: User = Depends(get_current_user), db=Depends(get_db)):
    """Natural language AI Command Center orchestrating multi-action tool calls."""
    history = [
        {"role": "user", "content": f"Command Center Prompt: '{req.command}'. Perform all relevant actions (create tasks, schedule daily plan, save reminders/notes) using available tools and report back a full summary."}
    ]
    def bound_tool_executor(tool_name, tool_args):
        return execute_ai_tool(tool_name, tool_args, user_id=user.id)

    ai_text = generate_response_with_tools(history, tool_executor_func=bound_tool_executor)
    return {"command": req.command, "summary": ai_text}


@app.get("/api/health")
def health():
    api_key = os.getenv("GEMINI_API_KEY", "").strip().strip("'\"")
    ai_configured = bool(api_key and api_key != "your_api_key_here")
    return {"status": "ok", "version": "3.0.0", "aiConfigured": ai_configured}


@app.get("/api/ai/test")
def test_ai():
    api_key = os.getenv("GEMINI_API_KEY", "").strip().strip("'\"")
    if not api_key or api_key == "your_api_key_here":
        raise HTTPException(status_code=400, detail="GEMINI_API_KEY is missing or unconfigured in backend/.env")

    test_history = [{"role": "user", "content": "Hello, please confirm in 1 sentence that Nexus AI is operational."}]
    res = generate_response_with_tools(test_history)
    return {"status": "success", "prompt": "Hello", "response": res}
